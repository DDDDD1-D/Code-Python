#encoding:UTF-8
import re
import requests
from bs4 import BeautifulSoup

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

import time

import string

from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication  
import smtplib

class auto_weekly_report:
	def __init__(self, url_abstract, location_docx):
		self.url_abstract = url_abstract
		self.location_docx = location_docx

	def get_abstract(self):
		page_content = BeautifulSoup(requests.get(self.url_abstract).content, "lxml")
		abstract = str(page_content.find(attrs={"class":"abstractSection abstractInFull"}).contents[0].contents[0])
		return abstract

	def write_to_docx(self, myname, progress, abstract, next_week):
		document = Document()

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		run = p.add_run('Weekly Report')
		font = run.font
		font.name = "Times New Roman"
		font.bold = True
		font.size = Pt(16)

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.LEFT
		run = p.add_run('Name: ' + myname + ' '*80 + 'Date: '+time.strftime('%Y-%m-%d',time.localtime(time.time())))
		font = run.font
		font.name = "Times New Roman"
		font.size = Pt(12)

		table = document.add_table(rows=3, cols=1)
		table.style = 'TableGrid'
		table.alignment = WD_TABLE_ALIGNMENT.LEFT
		table_cells = table.cell(0,0)
		table_cells.text = 'Progress:\n'
		p = table_cells.add_paragraph(progress)

		table_cells = table.cell(1,0)
		table_cells.text = 'Reading:\n' 
		p = table_cells.add_paragraph(abstract)

		table_cells = table.cell(2,0)
		table_cells.text = 'Next Week:\n' 
		p = table_cells.add_paragraph(next_week)
		
		document.save(self.location_docx)
		return "Successfully write the docx"

	def send_to_prof(self, address):
		msg = MIMEMultipart()

		msg.attach(MIMEText(''.join(open('./email_template').readlines())))

		attf = MIMEApplication(open(self.location_docx,'rb').read())  
		attf.add_header('Content-Disposition', 'attachment', filename='QQF'+time.strftime('%Y-%m-%d',time.localtime(time.time()))+'.docx')  
		msg.attach(attf)

		msgto = string.splitfields(address, ";") 
		msg['from'] = 'qqf14@mails.tsinghua.edu.cn'
		msg['subject'] = 'weekly report'

		try:
			server = smtplib.SMTP()
			server.connect('mails.tsinghua.edu.cn')
			server.login('qqf14@mails.tsinghua.edu.cn','XXXXX')
			server.sendmail(msg['from'], msgto, msg.as_string())
			server.quit()
			print "Successfully send to Profs"
		except Exception, e:
			print str(e)
		